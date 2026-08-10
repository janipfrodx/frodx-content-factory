# -*- coding: utf-8 -*-
"""
FrodX Newsletter — docx builder (GameChanger), shema = Janijev delujoči uvoznik.

KAKO UPORABIM (Claude):
  1. Preberi references/docx-pipeline.md (avtoritativna shema).
  2. V spodnji EDITIONS vpiši VSEBINO te izdaje za vse tri jezike (si/hr/en).
     - body razbij na BODY_P odstavke (seznam),
     - po želji dodaj bullets (seznam),
     - webinar blok dobi "event": [("EVENT_DATE","..."),("EVENT_TIME","..."),("EVENT_DURATION_MIN","...")],
     - sekundarna povezava (npr. video) = CTA bloka (inline povezav NI).
  3. Slike (PNG/JPG, ki jih je naložil Igor) daj v isto mapo; "img" = pot do datoteke.
     WebP najprej pretvori v PNG (PIL).
  4. Zaženi:  python3 scripts/build_newsletter.py
  5. Validiraj vsako datoteko z validate.py in poženi eval_check.py na docx.

ŽELEZNA PRAVILA: VELIKI ključi · hook/TOC/closing/signoff/PS so tabele ·
slika VDELANA v celico IMAGE · ena CTA na blok · vrstni red tabel fiksen.
"""
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from PIL import Image

NBSP = "\u00a0"; LQ = "\u201c"; RQ = "\u201d"  # EN narekovaji; SI »…« HR »…«

def pct(s):
    """NBSP pred % (samo SI/HR). EN ostane brez presledka."""
    return s.replace(" %", NBSP + "%") if isinstance(s, str) else s

# =====================================================================
# ============  VSEBINA IZDAJE — Claude to ZAMENJA  ===================
# =====================================================================
# Spodaj je PREDLOGA z vsemi polji (prikazan si; ponovi za hr in en).
# Vrednosti so vzorčne — prepiši jih z dejansko vsebino izdaje.
# Pomišljaj v telesu = en dash "–" (NE em dash "—").

EDITIONS = {
 "si": {
   "meta": [
     ("PACKAGE_ID","nl-YYYY-MM-slug"),
     ("EDITION_NAME","Tematski naslov izdaje."),
     ("LANGUAGE","si"),
     ("STATUS","ready_to_send"),                      # ali draft
     ("SEGMENT_REF","SI_ALL — celotna slovenska newsletter baza (Jani mapira na HubSpot list ID)"),
     ("SEND_DATETIME","2026-01-01T08:00:00"),         # Igor potrdi; isti čas čez vse jezike
     ("TIMEZONE","Europe/Ljubljana"),                 # HR = Europe/Zagreb
     ("FROM_NAME","Igor Pauletič"),                   # HR = Igor Pauletić
     ("FROM_EMAIL","igor.pauletic@frodx.com"),
     ("REPLY_TO","igor.pauletic@frodx.com"),
     ("FOOTER_REF","FRODX_SI"),                       # FRODX_HR / FRODX_EN
     ("SUBJECT","Subject po arhetipu."),
     ("PREHEADER","Preheader, ki nadaljuje subject."),
     ("GREETING","Pozdravljeni,"),                    # HR Pozdrav, · EN Hello,
   ],
   "hook_archetype":"B_prizor",                       # A_statistika/B_prizor/C_kontrast/D_univerzalna
   "hook":[
     "prvi odstavek hooka z malo začetnico (teče iz GREETING).",
     "drugi odstavek hooka.",
     "tretji odstavek hooka (po potrebi).",
   ],
   "toc":[
     ("Naslov v kazalu — blok 1","column","block-01"),
     # ("Naslov v kazalu — blok 2","webinar","block-02"),
   ],
   "blocks":[
     {"id":"block-01","type":"column","img":"block-01.png","img_file":"block-01.png",
      "img_alt":"Alt besedilo slike.",
      "title":"Naslov bloka 1.",
      "body":[
        "Prvi odstavek telesa.",
        "Drugi odstavek telesa. Pomišljaj kot en dash – takole.",
      ],
      "bullets":[                                      # po želji; sicer izpusti ali daj []
        "Prva alineja?",
        "Druga alineja?",
        "Tretja alineja?",
      ],
      "cta_label":"Preberi kolumno","cta_url":"https://frodx.com/sl/blog/..."},
     # --- PRIMER webinar bloka (event polja med bullets in CTA): ---
     # {"id":"block-02","type":"webinar","img":"block-02.png","img_file":"block-02.png",
     #  "img_alt":"Napoved webinarja.",
     #  "title":"Naslov webinarja.",
     #  "body":["Opis webinarja.","Kdo predava."],
     #  "bullets":["Točka 1","Točka 2"],
     #  "event":[("EVENT_DATE","2026-06-18"),("EVENT_TIME","10:00"),("EVENT_DURATION_MIN","45")],
     #  "cta_label":"Prijavite se na webinar","cta_url":"https://frodx.com/..."},
   ],
   "closing_type":"bookend",
   "closing":[
     "Navezava na hook (zaokrožitev zgodbe).",
     "Nizkofrikcijski odgovor-z-geslom: odgovorite na ta mail z »… check«.",
   ],
   "signoff_phrase":"Bodite dobro,","signoff_name":"Igor",
   "ps":"P.S. en odstavek (časovno omejen ali dopolnilni nagovor).",
 },
 # "hr": { ... ista struktura, nativna transkreacija, TIMEZONE Europe/Zagreb,
 #          FROM_NAME Igor Pauletić, GREETING "Pozdrav,", SIGNOFF "Sve najbolje," ... },
 # "en": { ... mednarodna poslovna angleščina, GREETING "Hello,", SIGNOFF "Best,",
 #          % brez presledka, narekovaji "…" ... },
}

OUT_DIR = "/mnt/user-data/outputs"

# =====================================================================
# ====================  ENGINE — ne spreminjaj  =======================
# =====================================================================
def kv_table(doc, rows):
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
    for k, v in rows:
        c = t.add_row().cells
        c[0].text = k
        if v != "__IMG__":
            c[1].text = v
    return t

def add_image_to_cell(cell, path):
    natural = Image.open(path).size[0]
    width_in = min(natural, 520) / 96.0          # skaliraj samo navzdol, razmerje ohrani
    cell.paragraphs[0].add_run().add_picture(path, width=Inches(width_in))

def build(lang, ed, outpath):
    P = pct if lang in ("si", "hr") else (lambda x: x)
    doc = Document()
    # python-docx pusti prazen <w:zoom> brez percent → pade na strogi validaciji; odstrani ga
    st = doc.settings.element; z = st.find(qn('w:zoom'))
    if z is not None: st.remove(z)
    doc.add_heading(f"FrodX Newsletter ({lang.upper()})", level=1)

    kv_table(doc, [(k, P(v)) for k, v in ed["meta"]]); doc.add_paragraph("")
    kv_table(doc, [("HOOK_ARCHETYPE", ed["hook_archetype"])] +
                  [(f"HOOK_P{i+1}", P(p)) for i, p in enumerate(ed["hook"])]); doc.add_paragraph("")

    t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"
    h = t.rows[0].cells; h[0].text, h[1].text, h[2].text = "TOC_LABEL", "TOC_TYPE", "BLOCK_ID"
    for label, typ, bid in ed["toc"]:
        c = t.add_row().cells; c[0].text, c[1].text, c[2].text = P(label), typ, bid
    doc.add_paragraph("")

    for b in ed["blocks"]:
        rows = [("BLOCK_ID", b["id"]), ("BLOCK_TYPE", b["type"]),
                ("IMAGE_FILE", b["img_file"]), ("IMAGE_ALT", P(b["img_alt"])),
                ("IMAGE", "__IMG__"), ("BLOCK_TITLE", P(b["title"]))]
        for i, p in enumerate(b["body"]):            rows.append((f"BODY_P{i+1}", P(p)))
        for i, bl in enumerate(b.get("bullets", [])): rows.append((f"BULLET_{i+1}", P(bl)))
        for k, v in b.get("event", []):              rows.append((k, v))
        rows += [("CTA_LABEL", P(b["cta_label"])), ("CTA_URL", b["cta_url"])]
        tbl = kv_table(doc, rows)
        for r in tbl.rows:
            if r.cells[0].text == "IMAGE":
                add_image_to_cell(r.cells[1], b["img"]); break
        doc.add_paragraph("")

    kv_table(doc, [("CLOSING_TYPE", ed["closing_type"])] +
                  [(f"CLOSING_P{i+1}", P(p)) for i, p in enumerate(ed["closing"])]); doc.add_paragraph("")
    kv_table(doc, [("SIGNOFF_PHRASE", ed["signoff_phrase"]), ("SIGNOFF_NAME", ed["signoff_name"])]); doc.add_paragraph("")
    kv_table(doc, [("PS_TEXT", P(ed["ps"]))])
    doc.save(outpath)
    print("zgrajeno:", outpath)

if __name__ == "__main__":
    import os
    os.makedirs(OUT_DIR, exist_ok=True)
    for lang, ed in EDITIONS.items():
        build(lang, ed, f"{OUT_DIR}/GameChanger_{lang}.docx")
